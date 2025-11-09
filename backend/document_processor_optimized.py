import logging
from pathlib import Path
from typing import List, Tuple
import json
import csv
from io import StringIO
from concurrent.futures import ThreadPoolExecutor
import multiprocessing

# PDF processing
import pdfplumber
from pdf2image import convert_from_path
import pytesseract
from PIL import Image

# Word processing
from docx import Document as DocxDocument

# Excel processing
from openpyxl import load_workbook

# ODT processing
from odf import text as odf_text, teletype
from odf.opendocument import load as odf_load

logger = logging.getLogger(__name__)


class OptimizedDocumentProcessor:
    """ULTRA-OPTIMIZED processor with semantic chunking and entity extraction"""
    
    def __init__(self, chunk_size: int = 400, chunk_overlap: int = 200, max_workers: int = None):
        # OPTIMIZED chunk settings for precision
        self.chunk_size = chunk_size  # Reduced from 800 for precision
        self.chunk_overlap = chunk_overlap  # Increased from 150 for context
        self.min_chunk_size = 50  # Reduced for finer granularity
        
        # Parallel processing
        if max_workers is None:
            max_workers = max(1, multiprocessing.cpu_count() - 1)
        self.max_workers = max_workers
        
        # Entity extractor (disabled by default for memory - patterns still work)
        try:
            from entity_extractor import EntityExtractor
            self.entity_extractor = EntityExtractor(enable_ner=False)  # Regex only, no NER
            logger.info("Entity extractor loaded (regex mode, NER disabled for memory)")
        except Exception as e:
            logger.warning(f"Could not load entity extractor: {e}")
            self.entity_extractor = None
        
        logger.info(f"ULTRA-OPTIMIZED processor: chunk_size={chunk_size}, overlap={chunk_overlap}, workers={max_workers}")
    
    def process_document(self, file_path: str) -> List[str]:
        """Process a document and return text chunks (optimized)"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        try:
            if extension == '.pdf':
                text = self._process_pdf_optimized(file_path)
            elif extension in ['.docx', '.doc']:
                text = self._process_word(file_path)
            elif extension in ['.xlsx', '.xls']:
                text = self._process_excel_optimized(file_path)
            elif extension == '.odt':
                text = self._process_odt(file_path)
            elif extension in ['.txt', '.md']:
                text = self._process_text(file_path)
            elif extension == '.json':
                text = self._process_json(file_path)
            elif extension == '.csv':
                text = self._process_csv(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return []
            
            # Split into chunks
            chunks = self._split_into_chunks(text)
            return chunks
        
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return []
    
    def _process_pdf_page(self, args: Tuple[str, int]) -> Tuple[int, str]:
        """Process a single PDF page (for parallel processing)"""
        file_path, page_num = args
        text = ""
        
        try:
            with pdfplumber.open(file_path) as pdf:
                if page_num < len(pdf.pages):
                    page = pdf.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text = page_text
        except Exception as e:
            logger.error(f"Error processing page {page_num} of {file_path}: {e}")
        
        return (page_num, text)
    
    def _process_pdf_page_ocr(self, args: Tuple[str, int]) -> Tuple[int, str]:
        """Process a single PDF page with OCR (for parallel processing)"""
        file_path, page_num = args
        text = ""
        
        try:
            # Convert single page to image
            images = convert_from_path(
                file_path,
                first_page=page_num + 1,
                last_page=page_num + 1,
                dpi=200  # Lower DPI for faster processing, still good quality
            )
            
            if images:
                # Optimized OCR settings for speed
                ocr_text = pytesseract.image_to_string(
                    images[0],
                    lang='eng+fra',
                    config='--psm 3 --oem 1'  # Faster OCR mode
                )
                text = ocr_text
        except Exception as e:
            logger.error(f"Error OCR processing page {page_num} of {file_path}: {e}")
        
        return (page_num, text)
    
    def _process_pdf_optimized(self, file_path: str) -> str:
        """Extract text from PDF with parallel processing and optimized OCR"""
        text = ""
        
        try:
            # First, try to extract text directly with parallel processing
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                if page_count == 0:
                    return ""
                
                # For small PDFs, process sequentially
                if page_count <= 3:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                else:
                    # For larger PDFs, use parallel processing
                    with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                        page_args = [(file_path, i) for i in range(page_count)]
                        results = list(executor.map(self._process_pdf_page, page_args))
                    
                    # Sort by page number and combine
                    results.sort(key=lambda x: x[0])
                    for page_num, page_text in results:
                        if page_text:
                            text += page_text + "\n"
            
            # If no text extracted, use parallel OCR
            if not text.strip():
                logger.info(f"No text found in PDF, using parallel OCR: {file_path}")
                
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    
                    # Limit OCR to reasonable number of pages for performance
                    max_ocr_pages = min(page_count, 50)  # Limit OCR to 50 pages
                    
                    if max_ocr_pages <= 3:
                        # Sequential OCR for small documents
                        images = convert_from_path(file_path, dpi=200)
                        for i, image in enumerate(images[:max_ocr_pages]):
                            ocr_text = pytesseract.image_to_string(
                                image,
                                lang='eng+fra',
                                config='--psm 3 --oem 1'
                            )
                            text += f"\n--- Page {i+1} ---\n{ocr_text}"
                    else:
                        # Parallel OCR for larger documents
                        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                            page_args = [(file_path, i) for i in range(max_ocr_pages)]
                            results = list(executor.map(self._process_pdf_page_ocr, page_args))
                        
                        # Sort by page number and combine
                        results.sort(key=lambda x: x[0])
                        for page_num, page_text in results:
                            if page_text:
                                text += f"\n--- Page {page_num+1} ---\n{page_text}"
                    
                    if page_count > max_ocr_pages:
                        logger.warning(f"PDF has {page_count} pages, only processed first {max_ocr_pages} with OCR")
        
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
        
        return text
    
    def _process_word(self, file_path: str) -> str:
        """Extract text from Word document (kept original for compatibility)"""
        text = ""
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
        
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
        
        return text
    
    def _process_excel_optimized(self, file_path: str) -> str:
        """Extract text from Excel file (optimized with read-only mode)"""
        text = ""
        
        try:
            # Use read_only and data_only for faster processing
            wb = load_workbook(file_path, read_only=True, data_only=True)
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                
                # Process rows in batches for better performance
                rows = []
                for row in sheet.iter_rows(values_only=True, max_row=10000):  # Limit rows for performance
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        rows.append(row_text)
                
                text += "\n".join(rows) + "\n"
        
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {e}")
        
        return text
    
    def _process_odt(self, file_path: str) -> str:
        """Extract text from ODT file"""
        text = ""
        
        try:
            doc = odf_load(file_path)
            all_paras = doc.getElementsByType(odf_text.P)
            
            for para in all_paras:
                para_text = teletype.extractText(para)
                text += para_text + "\n"
        
        except Exception as e:
            logger.error(f"Error processing ODT file {file_path}: {e}")
        
        return text
    
    def _process_text(self, file_path: str) -> str:
        """Extract text from plain text or markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return ""
    
    def _process_json(self, file_path: str) -> str:
        """Extract text from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error processing JSON file {file_path}: {e}")
            return ""
    
    def _process_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        text = ""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                csv_reader = csv.reader(f)
                for row in csv_reader:
                    text += " | ".join(row) + "\n"
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
        
        return text
    
    def _preprocess_text(self, text: str) -> str:
        """Preprocess text for better chunking and indexing"""
        if not text:
            return ""
        
        # Remove excessive whitespace while preserving structure
        lines = text.split('\n')
        processed_lines = []
        
        for line in lines:
            # Remove excessive spaces within lines
            cleaned_line = ' '.join(line.split())
            if cleaned_line:  # Only keep non-empty lines
                processed_lines.append(cleaned_line)
        
        # Join lines with single newlines
        processed_text = '\n'.join(processed_lines)
        
        # Remove excessive consecutive newlines (keep max 2)
        while '\n\n\n' in processed_text:
            processed_text = processed_text.replace('\n\n\n', '\n\n')
        
        return processed_text.strip()
    
    def _split_into_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks with intelligent boundary detection (optimized)"""
        if not text.strip():
            return []
        
        # Preprocess text
        text = self._preprocess_text(text)
        
        if len(text) < self.min_chunk_size:
            return [text] if text.strip() else []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + self.chunk_size
            
            # If not at the end, try to break at natural boundaries
            if end < text_length:
                # Priority 1: Paragraph boundary (double newline)
                para_break = text.rfind('\n\n', start, end)
                if para_break != -1 and para_break > start + self.chunk_size // 2:
                    end = para_break + 2
                else:
                    # Priority 2: Sentence boundary
                    sentence_delimiters = ['. ', '.\n', '! ', '!\n', '? ', '?\n', '.\t']
                    best_delim_pos = -1
                    
                    for delimiter in sentence_delimiters:
                        search_start = start + self.chunk_size // 2
                        delim_pos = text.rfind(delimiter, search_start, end)
                        if delim_pos > best_delim_pos:
                            best_delim_pos = delim_pos
                    
                    if best_delim_pos != -1:
                        for delimiter in sentence_delimiters:
                            if text[best_delim_pos:best_delim_pos+len(delimiter)] == delimiter:
                                end = best_delim_pos + len(delimiter)
                                break
                    else:
                        # Priority 3: Line boundary
                        line_break = text.rfind('\n', start + self.chunk_size // 2, end)
                        if line_break != -1:
                            end = line_break + 1
                        else:
                            # Priority 4: Word boundary
                            last_space = text.rfind(' ', start + self.chunk_size // 2, end)
                            if last_space != -1:
                                end = last_space + 1
            
            # Extract chunk
            chunk = text[start:end].strip()
            
            # Only add chunks that meet minimum size
            if len(chunk) >= self.min_chunk_size:
                chunks.append(chunk)
            elif len(chunk) > 0 and start + self.chunk_size >= text_length:
                chunks.append(chunk)
            
            # Move to next chunk with overlap
            if end >= text_length:
                break
            
            start = max(start + 1, end - self.chunk_overlap)
        
        return chunks
